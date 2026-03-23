#!/usr/bin/env python3
"""
Convert Markdown file to Word document with support for tables, bold text, and multi-level bullets
Enhanced custom implementation that actually works
"""
import sys
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import re
import os
import tempfile
import base64
import json
from io import BytesIO

# Try to import requests, install if missing
try:
    import requests
except ImportError:
    print("requests library not found. Please install it with: pip install requests")
    print("Mermaid diagrams will be rendered as text instead of images.")
    requests = None

def add_formatted_paragraph(doc, text, style=None):
    """Add a paragraph with bold and italic formatting and line breaks"""
    # Handle line breaks first - split by <br> or <BR> tags
    lines = re.split(r'<[Bb][Rr]\s*/?>', text)
    
    for line_idx, line in enumerate(lines):
        if line_idx > 0:
            # Add a new paragraph for each line after the first
            p = doc.add_paragraph(style=style)
        else:
            # First line uses the main paragraph
            p = doc.add_paragraph(style=style)
        
        # Split text by formatting markers
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Italic text
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Inline code
                run = p.add_run(part[1:-1])
                run.font.name = 'Courier New'
            else:
                # Regular text
                if part:
                    p.add_run(part)
    
    return p

def get_indent_level(line):
    """Get the indentation level of a line"""
    return (len(line) - len(line.lstrip())) // 2

def process_mermaid_block(lines, start_idx):
    """Process Mermaid code block and return content and next line index"""
    mermaid_lines = []
    i = start_idx + 1  # Skip the opening ```mermaid
    
    # Collect all lines until closing ```
    while i < len(lines):
        line = lines[i]
        if line.strip() == '```':
            break
        mermaid_lines.append(line)
        i += 1
    
    return '\n'.join(mermaid_lines), i + 1

def render_mermaid_to_image(mermaid_code):
    """Render Mermaid diagram to image using online service"""
    # Check if requests is available
    if requests is None:
        print("✗ requests library not available. Cannot render Mermaid as image.")
        return None
        
    try:
        print(f"Attempting to render Mermaid diagram with {len(mermaid_code)} characters...")
        
        # Clean up the mermaid code - remove excessive whitespace but preserve line breaks
        mermaid_code = '\n'.join(line.strip() for line in mermaid_code.split('\n') if line.strip())
        
        # Try different services for better compatibility
        services = [
            {
                'name': 'mermaid.ink',
                'url_template': 'https://mermaid.ink/img/{}',
                'encoding': 'base64'
            },
            {
                'name': 'kroki.io',
                'url_template': 'https://kroki.io/mermaid/png/{}',
                'encoding': 'base64'
            }
        ]
        
        for service in services:
            try:
                print(f"Trying {service['name']} service...")
                
                if service['encoding'] == 'base64':
                    # Use base64 encoding
                    import urllib.parse
                    encoded = base64.b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
                    # URL encode the base64 to handle special characters
                    encoded = urllib.parse.quote(encoded, safe='')
                else:
                    # Direct URL encoding
                    import urllib.parse
                    encoded = urllib.parse.quote(mermaid_code, safe='')
                
                url = service['url_template'].format(encoded)
                print(f"Requesting from: {url[:100]}...")
                
                # Download the image
                response = requests.get(url, timeout=30, headers={
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                })
                
                print(f"Response status: {response.status_code}")
                print(f"Response content type: {response.headers.get('content-type', 'unknown')}")
                print(f"Response content length: {len(response.content)} bytes")
                
                if response.status_code == 200:
                    # Check if we actually got an image
                    content_type = response.headers.get('content-type', '').lower()
                    if 'image' in content_type or len(response.content) > 1000:  # Assume images are > 1KB
                        print(f"✓ Successfully received image data from {service['name']}")
                        return BytesIO(response.content)
                    else:
                        print(f"✗ Response from {service['name']} doesn't appear to be an image. Content: {response.text[:200]}...")
                        continue
                else:
                    print(f"✗ {service['name']} failed with status code: {response.status_code}")
                    if response.status_code == 404:
                        print("Diagram might have syntax errors or be too complex")
                    print(f"Response: {response.text[:300]}...")
                    continue
                    
            except requests.RequestException as e:
                print(f"✗ Network error with {service['name']}: {e}")
                continue
            except Exception as e:
                print(f"✗ Error with {service['name']}: {e}")
                continue
        
        print("✗ All services failed to render Mermaid diagram")
        return None
            
    except Exception as e:
        print(f"✗ Critical error rendering Mermaid diagram: {e}")
        return None

def process_image(image_path, base_path=None):
    """Process an image (local or remote) and return a BytesIO stream"""
    try:
        # Check if it's a URL
        if image_path.startswith(('http://', 'https://')):
            # Download remote image
            response = requests.get(image_path, timeout=30)
            if response.status_code == 200:
                return BytesIO(response.content)
            else:
                print(f"Failed to download image from {image_path}. Status code: {response.status_code}")
                return None
        else:
            # Handle local file
            # Make path absolute if base_path is provided
            if base_path and not os.path.isabs(image_path):
                image_path = os.path.join(base_path, image_path)
            
            # Check if file exists
            if os.path.exists(image_path):
                with open(image_path, 'rb') as f:
                    return BytesIO(f.read())
            else:
                print(f"Image file not found: {image_path}")
                return None
                
    except Exception as e:
        print(f"Error processing image {image_path}: {e}")
        return None

def add_image_to_document(doc, image_stream, alt_text="", max_width=6.0):
    """Add an image to the document with proper sizing"""
    try:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        
        # Add image with max width constraint
        run.add_picture(image_stream, width=Inches(max_width))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add alt text as caption if provided
        if alt_text:
            caption = doc.add_paragraph()
            caption_run = caption.add_run(alt_text)
            caption_run.italic = True
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run.font.color.rgb = RGBColor(96, 96, 96)  # Gray
        
        return True
        
    except Exception as e:
        print(f"Failed to add image to document: {e}")
        return False

def process_table(lines, start_idx):
    """Process markdown table and return table object and next line index"""
    table_lines = []
    i = start_idx
    
    # Collect all table lines
    while i < len(lines):
        line = lines[i].strip()
        if '|' in line:
            table_lines.append(line)
            i += 1
        else:
            break
    
    if len(table_lines) < 2:
        return None, start_idx
    
    # Parse table
    rows = []
    for line in table_lines:
        if line.startswith('|') and line.endswith('|'):
            line = line[1:-1]  # Remove outer pipes
        cells = [cell.strip() for cell in line.split('|')]
        # Skip separator row (contains only dashes and spaces)
        if not re.match(r'^[\s\-\|:]+$', line):
            rows.append(cells)
    
    return rows, i

def markdown_to_word(md_file_path, output_path):
    """Convert markdown file to Word document"""
    
    # Read the markdown file
    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
    except FileNotFoundError:
        print(f"Error: File '{md_file_path}' not found.")
        return False
    except Exception as e:
        print(f"Error reading file: {e}")
        return False
    
    # Create a new Word document
    doc = Document()
    
    # Get base path for relative image paths
    base_path = os.path.dirname(os.path.abspath(md_file_path))
    
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    numbered_list_counter = {}  # Track numbering for each indent level
    last_was_numbered_list = False
    
    while i < len(lines):
        line = lines[i]
        original_line = line
        line = line.rstrip()
        
        # Handle code blocks
        if line.startswith('```'):
            if not in_code_block:
                # Check if this is a Mermaid diagram
                if line.strip().lower() in ['```mermaid', '```mermaid '] or line.strip().startswith('```mermaid'):
                    mermaid_content, next_i = process_mermaid_block(lines, i)
                    
                    print(f"Found Mermaid diagram:")
                    print(f"First 100 chars: {mermaid_content[:100]}...")
                    
                    # Try to render Mermaid diagram as image
                    image_stream = render_mermaid_to_image(mermaid_content)
                    
                    if image_stream:
                        # Add the rendered image to the document
                        paragraph = doc.add_paragraph()
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        try:
                            # Add image with reasonable size (max 6 inches wide)
                            run.add_picture(image_stream, width=Inches(6))
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            print("✓ Successfully added Mermaid image to document")
                        except Exception as e:
                            print(f"✗ Failed to add image to document: {e}")
                            # Fallback to text representation
                            doc.add_paragraph().add_run('Mermaid Diagram (failed to render as image):').bold = True
                            p = doc.add_paragraph()
                            p.style = 'List Paragraph'
                            run = p.add_run(mermaid_content)
                            run.font.name = 'Courier New'
                    else:
                        # Fallback to text representation if rendering fails
                        print("Using fallback text representation")
                        heading = doc.add_paragraph()
                        heading.add_run('Mermaid Diagram (rendered as text):').bold = True
                        
                        # Add the Mermaid code in a code block style
                        p = doc.add_paragraph()
                        p.style = 'List Paragraph'
                        run = p.add_run(mermaid_content)
                        run.font.name = 'Courier New'
                        run.font.color.rgb = RGBColor(64, 64, 64)  # Dark gray
                        
                        # Add note about failed rendering
                        note = doc.add_paragraph()
                        note_run = note.add_run('Note: Failed to render Mermaid diagram as image. Displaying source code.')
                        note_run.italic = True
                        note_run.font.color.rgb = RGBColor(255, 0, 0)  # Red for error
                    
                    i = next_i
                    continue
                else:
                    in_code_block = True
            else:
                in_code_block = False
            i += 1
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'List Paragraph'
            for run in p.runs:
                run.font.name = 'Courier New'
            i += 1
            continue
        
        if not line.strip():
            # Empty line - reset numbered list tracking
            if last_was_numbered_list:
                numbered_list_counter = {}
                last_was_numbered_list = False
            i += 1
            continue
        
        # Check for tables
        if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_rows, next_i = process_table(lines, i)
            if table_rows and len(table_rows) > 0:
                # Create table
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Table Grid'
                
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(table.rows[row_idx].cells):
                            cell = table.rows[row_idx].cells[col_idx]
                            # Add formatted text to cell
                            cell.text = ''
                            p = cell.paragraphs[0]
                            p.clear()
                            
                            # Apply formatting within cell (including line breaks)
                            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', cell_data)
                            
                            # Handle line breaks in table cells
                            cell_lines = re.split(r'<[Bb][Rr]\s*/?>', cell_data)
                            
                            for line_idx, cell_line in enumerate(cell_lines):
                                if line_idx > 0:
                                    # Add a line break in the cell
                                    p.add_run('\n')
                                
                                # Apply formatting to this line
                                line_parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', cell_line)
                                for part in line_parts:
                                    if part.startswith('**') and part.endswith('**'):
                                        run = p.add_run(part[2:-2])
                                        run.bold = True
                                    elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                                        run = p.add_run(part[1:-1])
                                        run.italic = True
                                    elif part.startswith('`') and part.endswith('`'):
                                        run = p.add_run(part[1:-1])
                                        run.font.name = 'Courier New'
                                    else:
                                        if part:
                                            p.add_run(part)
                            
                            # Bold header row
                            if row_idx == 0:
                                for run in p.runs:
                                    run.bold = True
                
                i = next_i
                continue
        
        # Handle headers
        # Map Markdown heading levels to Word styles:
        # H1 -> Title (level=0), H2 -> Heading 1 (level=1), H3 -> Heading 2 (level=2), etc.
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=0)
            # Reset numbering when encountering headers
            numbered_list_counter = {}
            last_was_numbered_list = False
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=1)
            numbered_list_counter = {}
            last_was_numbered_list = False
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=2)
            numbered_list_counter = {}
            last_was_numbered_list = False
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=3)
            numbered_list_counter = {}
            last_was_numbered_list = False
        elif line.startswith('##### '):
            heading = doc.add_heading(line[6:], level=4)
            numbered_list_counter = {}
            last_was_numbered_list = False
        elif line.startswith('###### '):
            heading = doc.add_heading(line[7:], level=5)
            numbered_list_counter = {}
            last_was_numbered_list = False
        # Handle lists with multiple levels
        elif re.match(r'^(\s*)([-*]|\d+\.)\s', line):
            indent_level = get_indent_level(original_line)
            
            # Check if this is a numbered list item
            is_numbered = re.match(r'^\s*\d+\.\s', line)
            
            # Extract list content
            if re.match(r'^\s*[-*]\s', line):
                content = re.sub(r'^\s*[-*]\s*', '', line)
                if indent_level == 0:
                    style = 'List Bullet'
                elif indent_level == 1:
                    style = 'List Bullet 2'
                elif indent_level == 2:
                    style = 'List Bullet 3'
                else:
                    style = 'List Paragraph'
                # Don't reset numbered list tracking for bullet lists - they can be mixed
                last_was_numbered_list = False
                # Add formatted paragraph for bullet lists
                add_formatted_paragraph(doc, content, style=style)
            elif is_numbered:
                content = re.sub(r'^\s*\d+\.\s*', '', line)
                
                # Reset all deeper levels when returning to a shallower level
                keys_to_remove = [k for k in numbered_list_counter.keys() if k > indent_level]
                for k in keys_to_remove:
                    del numbered_list_counter[k]
                
                # Check if we need to initialize or increment numbering for this level
                if indent_level not in numbered_list_counter:
                    numbered_list_counter[indent_level] = 1
                else:
                    numbered_list_counter[indent_level] += 1
                
                # Create custom numbered paragraph with explicit numbering
                if indent_level == 0:
                    style = 'Normal'
                    number_prefix = f"{numbered_list_counter[indent_level]}. "
                elif indent_level == 1:
                    style = 'List Paragraph'
                    number_prefix = f"  {numbered_list_counter[indent_level]}. "
                elif indent_level == 2:
                    style = 'List Paragraph'
                    number_prefix = f"    {numbered_list_counter[indent_level]}. "
                else:
                    style = 'List Paragraph'
                    number_prefix = "  " * indent_level + f"{numbered_list_counter[indent_level]}. "
                
                # Create paragraph and add number prefix without formatting, then formatted content
                p = doc.add_paragraph(style=style)
                # Add the number prefix (not formatted)
                p.add_run(number_prefix)
                
                # Add the content with formatting
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', content)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                        # Italic text
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        # Inline code
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                    else:
                        # Regular text
                        if part:
                            p.add_run(part)
                
                last_was_numbered_list = True
            else:
                content = line.strip()
                style = 'List Paragraph'
                # Don't automatically reset for other list types
                last_was_numbered_list = False
                # Add formatted paragraph for non-numbered lists
                add_formatted_paragraph(doc, content, style=style)
        # Handle regular paragraphs
        else:
            # Check for markdown images before processing as regular text
            image_match = re.search(r'!\[([^\]]*)\]\(([^\)]+)\)', line)
            if image_match:
                alt_text = image_match.group(1)
                image_path = image_match.group(2)
                
                # Process the text before the image (if any)
                text_before = line[:image_match.start()].strip()
                if text_before:
                    add_formatted_paragraph(doc, text_before)
                
                # Process the image
                image_stream = process_image(image_path, base_path)
                if image_stream:
                    if not add_image_to_document(doc, image_stream, alt_text):
                        # Fallback: add as text if image processing fails
                        fallback_text = f"[Image: {alt_text or image_path}]"
                        add_formatted_paragraph(doc, fallback_text)
                else:
                    # Add placeholder text for failed images
                    fallback_text = f"[Image not found: {alt_text or image_path}]"
                    p = add_formatted_paragraph(doc, fallback_text)
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Red for error
                
                # Process the text after the image (if any)
                text_after = line[image_match.end():].strip()
                if text_after:
                    add_formatted_paragraph(doc, text_after)
            else:
                # Regular paragraph without images
                add_formatted_paragraph(doc, line)
            
            # Reset numbered list tracking for regular paragraphs
            if last_was_numbered_list:
                numbered_list_counter = {}
            last_was_numbered_list = False
        
        i += 1
    
    # Save the document
    try:
        # Check if file exists and is potentially locked
        if os.path.exists(output_path):
            try:
                # Try to open the file to check if it's locked
                with open(output_path, 'r+b') as test_file:
                    pass  # File is not locked
            except (PermissionError, OSError) as e:
                print(f"Warning: Output file appears to be open in another application.")
                print(f"Please close '{output_path}' and try again.")
                return False
        
        doc.save(output_path)
        print(f"Successfully converted '{md_file_path}' to '{output_path}'")
        return True
    except PermissionError:
        print(f"Error: Permission denied saving to '{output_path}'")
        print("The file might be open in Word or another application. Please close it and try again.")
        return False
    except Exception as e:
        print(f"Error saving Word document: {e}")
        return False

if __name__ == "__main__":
    if len(sys.argv) < 2 or len(sys.argv) > 3:
        print("Usage: python md_to_word.py <input.md> [output.docx]")
        print("       python md_to_word.py <input.md>  (output will be same name with .docx extension)")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    if len(sys.argv) == 3:
        output_file = sys.argv[2]
    else:
        # Generate output filename from input filename
        if input_file.lower().endswith('.md'):
            output_file = input_file[:-3] + '.docx'
        else:
            output_file = input_file + '.docx'
    
    if markdown_to_word(input_file, output_file):
        print("Conversion completed successfully!")
    else:
        print("Conversion failed!")
        sys.exit(1)